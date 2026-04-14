from docx import Document

doc = Document()
doc.add_heading('PHOTOSYNTHESIS: A COMPREHENSIVE STUDY GUIDE', 0)

with open('c:\\Users\\NC\\Desktop\\ai-study-assistant\\test-document.txt', 'r') as f:
    content = f.read()

doc.add_paragraph(content)

doc.save('c:\\Users\\NC\\Desktop\\ai-study-assistant\\test-document.docx')